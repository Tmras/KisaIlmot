from docx import Document
from docxcompose.composer import Composer


def create_files(contests):
    with open('readme.txt', 'w') as f:
        for contest in contests:

            f.write(f"{contest.date}|{contest.contest}\n")

def create_files_docx(contests):
    """
    Create the context documentation by adding the found contest data to template file
    :param contests: Contest data found from web
    :return:
    """
    document = Document(".//Templates//kisa.docx")
    composer = Composer(document)

    i = 0
    while i < len(contests) - 1:
        template = Document(".//Templates//kisa.docx")
        contest_calc = 0

        # Fill the tables in templates
        for table in template.tables:
            ind = i + contest_calc
            # Find the correct place in the template and put the information in it
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "kisa" in p.text:
                            inline = p.runs
                            for k in range(len(inline)):
                                if "kisa" in inline[k].text:
                                    text = inline[k].text.replace("kisa", contests[ind].contest)
                                    inline[k].text = text
                        elif "pvm"  in p.text:
                            inline = p.runs
                            for j in range(len(inline)):
                                if "pvm" in inline[j].text:
                                    text = inline[j].text.replace("pvm", contests[ind].date)
                                    inline[j].text = text
            if contest_calc == 0:
                contest_calc = 1
            else:
                contest_calc = 0

        composer.append(template)
        i = i + 2

    composer.save("kisat2.docx")
